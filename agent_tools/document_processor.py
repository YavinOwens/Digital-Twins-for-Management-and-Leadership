"""
Document Processor for Digital Twins Management System

This module provides comprehensive document processing capabilities for various data types
including databases, structured, semi-structured, and unstructured datasets.
"""

import os
import pandas as pd
import json
import csv
import sqlite3
import logging
from typing import Dict, List, Any, Optional, Union
from pathlib import Path
import streamlit as st
from io import BytesIO, StringIO
import zipfile
import tarfile
import xml.etree.ElementTree as ET
import yaml
import pickle
import openpyxl
from docx import Document
from PyPDF2 import PdfReader
import chardet
import mimetypes

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Comprehensive document processor for various data types
    """
    
    def __init__(self, upload_dir: str = "./uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
        self.supported_formats = {
            # Structured data
            'csv': self._process_csv,
            'xlsx': self._process_excel,
            'xls': self._process_excel,
            'json': self._process_json,
            'xml': self._process_xml,
            'yaml': self._process_yaml,
            'yml': self._process_yaml,
            'sqlite': self._process_sqlite,
            'db': self._process_sqlite,
            
            # Semi-structured data
            'txt': self._process_text,
            'md': self._process_markdown,
            'log': self._process_log,
            'tsv': self._process_tsv,
            
            # Unstructured data
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'doc': self._process_docx,
            'rtf': self._process_rtf,
            
            # Archives
            'zip': self._process_archive,
            'tar': self._process_archive,
            'gz': self._process_archive,
        }
    
    def process_uploaded_file(self, uploaded_file) -> Dict[str, Any]:
        """
        Process an uploaded file and return structured data
        """
        try:
            # Save uploaded file
            file_path = self.upload_dir / uploaded_file.name
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            # Detect file type
            file_extension = Path(uploaded_file.name).suffix.lower().lstrip('.')
            mime_type, _ = mimetypes.guess_type(uploaded_file.name)
            
            # Process file
            if file_extension in self.supported_formats:
                processor = self.supported_formats[file_extension]
                result = processor(file_path)
            else:
                # Try to process as text
                result = self._process_text(file_path)
            
            # Add metadata
            result['metadata'] = {
                'filename': uploaded_file.name,
                'file_size': uploaded_file.size,
                'file_type': file_extension,
                'mime_type': mime_type,
                'upload_timestamp': pd.Timestamp.now().isoformat()
            }
            
            logger.info(f"Successfully processed file: {uploaded_file.name}")
            return result
            
        except Exception as e:
            logger.error(f"Error processing file {uploaded_file.name}: {e}")
            return {
                'error': str(e),
                'metadata': {
                    'filename': uploaded_file.name,
                    'file_size': uploaded_file.size,
                    'file_type': file_extension,
                    'mime_type': mime_type,
                    'upload_timestamp': pd.Timestamp.now().isoformat()
                }
            }
    
    def process_directory(self, directory_path: str) -> Dict[str, Any]:
        """
        Process all files in a directory
        """
        results = {}
        directory = Path(directory_path)
        
        if not directory.exists():
            return {'error': f"Directory {directory_path} does not exist"}
        
        for file_path in directory.rglob('*'):
            if file_path.is_file():
                try:
                    file_extension = file_path.suffix.lower().lstrip('.')
                    if file_extension in self.supported_formats:
                        processor = self.supported_formats[file_extension]
                        result = processor(file_path)
                        results[str(file_path)] = result
                except Exception as e:
                    logger.error(f"Error processing {file_path}: {e}")
                    results[str(file_path)] = {'error': str(e)}
        
        return results
    
    def _process_csv(self, file_path: Path) -> Dict[str, Any]:
        """Process CSV files"""
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                encoding = chardet.detect(raw_data)['encoding']
            
            # Read CSV
            df = pd.read_csv(file_path, encoding=encoding)
            
            return {
                'type': 'structured',
                'data_type': 'csv',
                'data': df.to_dict('records'),
                'columns': df.columns.tolist(),
                'shape': df.shape,
                'dtypes': df.dtypes.to_dict(),
                'summary': df.describe().to_dict() if len(df) > 0 else {},
                'null_counts': df.isnull().sum().to_dict(),
                'sample_data': df.head().to_dict('records')
            }
        except Exception as e:
            return {'error': f"CSV processing error: {e}"}
    
    def _process_excel(self, file_path: Path) -> Dict[str, Any]:
        """Process Excel files"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                sheets_data[sheet_name] = {
                    'data': df.to_dict('records'),
                    'columns': df.columns.tolist(),
                    'shape': df.shape,
                    'dtypes': df.dtypes.to_dict(),
                    'summary': df.describe().to_dict() if len(df) > 0 else {},
                    'null_counts': df.isnull().sum().to_dict(),
                    'sample_data': df.head().to_dict('records')
                }
            
            return {
                'type': 'structured',
                'data_type': 'excel',
                'sheets': sheets_data,
                'sheet_names': excel_file.sheet_names
            }
        except Exception as e:
            return {'error': f"Excel processing error: {e}"}
    
    def _process_json(self, file_path: Path) -> Dict[str, Any]:
        """Process JSON files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Determine if it's structured data (list of objects)
            if isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict):
                df = pd.DataFrame(data)
                return {
                    'type': 'structured',
                    'data_type': 'json',
                    'data': data,
                    'columns': df.columns.tolist(),
                    'shape': df.shape,
                    'dtypes': df.dtypes.to_dict(),
                    'summary': df.describe().to_dict() if len(df) > 0 else {},
                    'null_counts': df.isnull().sum().to_dict(),
                    'sample_data': df.head().to_dict('records')
                }
            else:
                return {
                    'type': 'semi_structured',
                    'data_type': 'json',
                    'data': data,
                    'structure': self._analyze_json_structure(data)
                }
        except Exception as e:
            return {'error': f"JSON processing error: {e}"}
    
    def _process_xml(self, file_path: Path) -> Dict[str, Any]:
        """Process XML files"""
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # Convert XML to dictionary
            def xml_to_dict(element):
                result = {}
                if element.text and element.text.strip():
                    result['text'] = element.text.strip()
                if element.attrib:
                    result['attributes'] = element.attrib
                for child in element:
                    child_data = xml_to_dict(child)
                    if child.tag in result:
                        if not isinstance(result[child.tag], list):
                            result[child.tag] = [result[child.tag]]
                        result[child.tag].append(child_data)
                    else:
                        result[child.tag] = child_data
                return result
            
            xml_data = xml_to_dict(root)
            
            return {
                'type': 'semi_structured',
                'data_type': 'xml',
                'data': xml_data,
                'root_tag': root.tag,
                'structure': self._analyze_xml_structure(root)
            }
        except Exception as e:
            return {'error': f"XML processing error: {e}"}
    
    def _process_yaml(self, file_path: Path) -> Dict[str, Any]:
        """Process YAML files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = yaml.safe_load(f)
            
            return {
                'type': 'semi_structured',
                'data_type': 'yaml',
                'data': data,
                'structure': self._analyze_yaml_structure(data)
            }
        except Exception as e:
            return {'error': f"YAML processing error: {e}"}
    
    def _process_sqlite(self, file_path: Path) -> Dict[str, Any]:
        """Process SQLite database files"""
        try:
            conn = sqlite3.connect(file_path)
            cursor = conn.cursor()
            
            # Get table names
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
            tables = [row[0] for row in cursor.fetchall()]
            
            tables_data = {}
            for table in tables:
                # Get table schema
                cursor.execute(f"PRAGMA table_info({table})")
                columns = [{'name': row[1], 'type': row[2], 'not_null': bool(row[3]), 'default': row[4], 'pk': bool(row[5])} for row in cursor.fetchall()]
                
                # Get sample data
                cursor.execute(f"SELECT * FROM {table} LIMIT 100")
                sample_data = [dict(zip([col[0] for col in cursor.description], row)) for row in cursor.fetchall()]
                
                # Get row count
                cursor.execute(f"SELECT COUNT(*) FROM {table}")
                row_count = cursor.fetchone()[0]
                
                tables_data[table] = {
                    'columns': columns,
                    'sample_data': sample_data,
                    'row_count': row_count
                }
            
            conn.close()
            
            return {
                'type': 'structured',
                'data_type': 'sqlite',
                'tables': tables_data,
                'table_names': tables
            }
        except Exception as e:
            return {'error': f"SQLite processing error: {e}"}
    
    def _process_text(self, file_path: Path) -> Dict[str, Any]:
        """Process text files"""
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                encoding = chardet.detect(raw_data)['encoding']
            
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            
            # Basic text analysis
            lines = content.split('\n')
            words = content.split()
            
            return {
                'type': 'unstructured',
                'data_type': 'text',
                'content': content,
                'line_count': len(lines),
                'word_count': len(words),
                'char_count': len(content),
                'encoding': encoding,
                'sample_lines': lines[:10] if len(lines) > 10 else lines
            }
        except Exception as e:
            return {'error': f"Text processing error: {e}"}
    
    def _process_markdown(self, file_path: Path) -> Dict[str, Any]:
        """Process Markdown files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Basic markdown analysis
            lines = content.split('\n')
            headers = [line for line in lines if line.strip().startswith('#')]
            
            return {
                'type': 'semi_structured',
                'data_type': 'markdown',
                'content': content,
                'line_count': len(lines),
                'word_count': len(content.split()),
                'char_count': len(content),
                'headers': headers,
                'sample_lines': lines[:10] if len(lines) > 10 else lines
            }
        except Exception as e:
            return {'error': f"Markdown processing error: {e}"}
    
    def _process_log(self, file_path: Path) -> Dict[str, Any]:
        """Process log files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            lines = content.split('\n')
            
            # Basic log analysis
            error_lines = [line for line in lines if 'error' in line.lower() or 'exception' in line.lower()]
            warning_lines = [line for line in lines if 'warning' in line.lower() or 'warn' in line.lower()]
            
            return {
                'type': 'semi_structured',
                'data_type': 'log',
                'content': content,
                'line_count': len(lines),
                'error_count': len(error_lines),
                'warning_count': len(warning_lines),
                'error_lines': error_lines[:10],  # First 10 errors
                'warning_lines': warning_lines[:10],  # First 10 warnings
                'sample_lines': lines[:20] if len(lines) > 20 else lines
            }
        except Exception as e:
            return {'error': f"Log processing error: {e}"}
    
    def _process_tsv(self, file_path: Path) -> Dict[str, Any]:
        """Process TSV files"""
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                encoding = chardet.detect(raw_data)['encoding']
            
            # Read TSV
            df = pd.read_csv(file_path, sep='\t', encoding=encoding)
            
            return {
                'type': 'structured',
                'data_type': 'tsv',
                'data': df.to_dict('records'),
                'columns': df.columns.tolist(),
                'shape': df.shape,
                'dtypes': df.dtypes.to_dict(),
                'summary': df.describe().to_dict() if len(df) > 0 else {},
                'null_counts': df.isnull().sum().to_dict(),
                'sample_data': df.head().to_dict('records')
            }
        except Exception as e:
            return {'error': f"TSV processing error: {e}"}
    
    def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Process PDF files"""
        try:
            reader = PdfReader(file_path)
            text_content = ""
            page_count = len(reader.pages)
            
            for page in reader.pages:
                text_content += page.extract_text() + "\n"
            
            return {
                'type': 'unstructured',
                'data_type': 'pdf',
                'content': text_content,
                'page_count': page_count,
                'word_count': len(text_content.split()),
                'char_count': len(text_content),
                'sample_text': text_content[:1000] if len(text_content) > 1000 else text_content
            }
        except Exception as e:
            return {'error': f"PDF processing error: {e}"}
    
    def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Process DOCX files"""
        try:
            doc = Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            return {
                'type': 'unstructured',
                'data_type': 'docx',
                'content': text_content,
                'paragraph_count': len(doc.paragraphs),
                'word_count': len(text_content.split()),
                'char_count': len(text_content),
                'sample_text': text_content[:1000] if len(text_content) > 1000 else text_content
            }
        except Exception as e:
            return {'error': f"DOCX processing error: {e}"}
    
    def _process_rtf(self, file_path: Path) -> Dict[str, Any]:
        """Process RTF files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Basic RTF processing (remove RTF formatting codes)
            import re
            clean_content = re.sub(r'\\[a-z]+\d*', '', content)
            clean_content = re.sub(r'[{}]', '', clean_content)
            
            return {
                'type': 'unstructured',
                'data_type': 'rtf',
                'content': clean_content,
                'word_count': len(clean_content.split()),
                'char_count': len(clean_content),
                'sample_text': clean_content[:1000] if len(clean_content) > 1000 else clean_content
            }
        except Exception as e:
            return {'error': f"RTF processing error: {e}"}
    
    def _process_archive(self, file_path: Path) -> Dict[str, Any]:
        """Process archive files (ZIP, TAR, etc.)"""
        try:
            archive_data = {}
            
            if file_path.suffix.lower() == '.zip':
                with zipfile.ZipFile(file_path, 'r') as zip_file:
                    for file_info in zip_file.filelist:
                        if not file_info.is_dir():
                            try:
                                content = zip_file.read(file_info.filename)
                                archive_data[file_info.filename] = {
                                    'size': file_info.file_size,
                                    'content': content.decode('utf-8', errors='ignore')[:1000]  # First 1000 chars
                                }
                            except:
                                archive_data[file_info.filename] = {
                                    'size': file_info.file_size,
                                    'content': '[Binary file]'
                                }
            
            elif file_path.suffix.lower() in ['.tar', '.gz']:
                with tarfile.open(file_path, 'r') as tar_file:
                    for member in tar_file.getmembers():
                        if member.isfile():
                            try:
                                content = tar_file.extractfile(member).read()
                                archive_data[member.name] = {
                                    'size': member.size,
                                    'content': content.decode('utf-8', errors='ignore')[:1000]  # First 1000 chars
                                }
                            except:
                                archive_data[member.name] = {
                                    'size': member.size,
                                    'content': '[Binary file]'
                                }
            
            return {
                'type': 'archive',
                'data_type': file_path.suffix.lower().lstrip('.'),
                'files': archive_data,
                'file_count': len(archive_data)
            }
        except Exception as e:
            return {'error': f"Archive processing error: {e}"}
    
    def _analyze_json_structure(self, data: Any, path: str = "") -> Dict[str, Any]:
        """Analyze JSON structure recursively"""
        if isinstance(data, dict):
            return {
                'type': 'object',
                'keys': list(data.keys()),
                'key_count': len(data.keys()),
                'children': {k: self._analyze_json_structure(v, f"{path}.{k}") for k, v in data.items()}
            }
        elif isinstance(data, list):
            return {
                'type': 'array',
                'length': len(data),
                'children': [self._analyze_json_structure(item, f"{path}[{i}]") for i, item in enumerate(data[:5])]  # First 5 items
            }
        else:
            return {
                'type': type(data).__name__,
                'value': str(data)[:100] if len(str(data)) > 100 else str(data)
            }
    
    def _analyze_xml_structure(self, element: ET.Element, path: str = "") -> Dict[str, Any]:
        """Analyze XML structure recursively"""
        children = {}
        for child in element:
            child_path = f"{path}.{child.tag}" if path else child.tag
            children[child.tag] = self._analyze_xml_structure(child, child_path)
        
        return {
            'tag': element.tag,
            'attributes': element.attrib,
            'text': element.text.strip() if element.text else None,
            'children': children,
            'child_count': len(children)
        }
    
    def _analyze_yaml_structure(self, data: Any, path: str = "") -> Dict[str, Any]:
        """Analyze YAML structure recursively"""
        if isinstance(data, dict):
            return {
                'type': 'mapping',
                'keys': list(data.keys()),
                'key_count': len(data.keys()),
                'children': {k: self._analyze_yaml_structure(v, f"{path}.{k}") for k, v in data.items()}
            }
        elif isinstance(data, list):
            return {
                'type': 'sequence',
                'length': len(data),
                'children': [self._analyze_yaml_structure(item, f"{path}[{i}]") for i, item in enumerate(data[:5])]  # First 5 items
            }
        else:
            return {
                'type': type(data).__name__,
                'value': str(data)[:100] if len(str(data)) > 100 else str(data)
            }
    
    def get_processed_data_summary(self, processed_data: Dict[str, Any]) -> str:
        """Generate a summary of processed data for agent context"""
        summary_parts = []
        
        if 'error' in processed_data:
            return f"Error processing file: {processed_data['error']}"
        
        metadata = processed_data.get('metadata', {})
        summary_parts.append(f"File: {metadata.get('filename', 'Unknown')}")
        summary_parts.append(f"Type: {processed_data.get('data_type', 'Unknown')}")
        summary_parts.append(f"Category: {processed_data.get('type', 'Unknown')}")
        
        if processed_data.get('type') == 'structured':
            if 'shape' in processed_data:
                summary_parts.append(f"Dimensions: {processed_data['shape']}")
            if 'columns' in processed_data:
                summary_parts.append(f"Columns: {', '.join(processed_data['columns'])}")
            if 'sheets' in processed_data:
                summary_parts.append(f"Sheets: {', '.join(processed_data['sheet_names'])}")
        
        elif processed_data.get('type') == 'semi_structured':
            if 'line_count' in processed_data:
                summary_parts.append(f"Lines: {processed_data['line_count']}")
            if 'headers' in processed_data:
                summary_parts.append(f"Headers: {len(processed_data['headers'])}")
        
        elif processed_data.get('type') == 'unstructured':
            if 'word_count' in processed_data:
                summary_parts.append(f"Words: {processed_data['word_count']}")
            if 'page_count' in processed_data:
                summary_parts.append(f"Pages: {processed_data['page_count']}")
        
        return " | ".join(summary_parts)
